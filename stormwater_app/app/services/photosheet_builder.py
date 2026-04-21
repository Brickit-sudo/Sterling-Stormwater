"""
app/services/photosheet_builder.py
DOCX builder for Photosheet mode — native Word header architecture.

Page structure:
  ┌─────────────────────────────────────────────────┐
  │ NATIVE WORD HEADER (auto-repeats every page):   │
  │   [Logo 2" | date / prepared_by right-aligned]  │
  │   Service photos are provided below.             │
  │   ─────────────────────────── (black rule)       │
  ├─────────────────────────────────────────────────┤
  │ BODY (single continuous table):                  │
  │   Each row: exact height = body_h / rows_per_pg │
  │   Word paginates naturally — zero manual breaks  │
  ├─────────────────────────────────────────────────┤
  │ NATIVE WORD FOOTER (auto-repeats every page):   │
  │   ─── separator, contact line, Page N of M      │
  └─────────────────────────────────────────────────┘

No manual page breaks in body.  No repeated header paragraphs.
No keepWithNext / keepLinesTogether / pageBreakBefore flags.
"""

import io
import math
import time
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from PIL import Image as PILImage, ImageOps

# ── Page geometry (inches) ────────────────────────────────────────────────────
_PAGE_W   = 8.5
_PAGE_H   = 11.0
_MARGIN_L = 0.60
_MARGIN_R = 0.60
# Top margin is computed dynamically from the actual logo height so the
# native header never overflows into the body area.  Computed in _init_geometry().
_MARGIN_T = 1.52   # default; overwritten by _init_geometry() at import time
_MARGIN_B = 0.85   # footer_distance (0.25") + contact + page line (~0.50")
_BODY_W   = _PAGE_W - _MARGIN_L - _MARGIN_R   # 7.30"
_BODY_H   = _PAGE_H - _MARGIN_T - _MARGIN_B   # recomputed by _init_geometry()

# Safety margin (inches) subtracted from body height before row allocation.
_ROW_SAFETY = 0.15

# Per-cell sizing
_DATE_H    = 0.18   # optional date line height (fixed — one line)
_CELL_PAD  = 0.04   # breathing room per cell

# Caption dynamic sizing (at 10pt Calibri bold)
_CAP_FONT_PT      = 10.0   # caption / service-photos font size
_CAP_LINE_H_PT    = 13.5   # line height including leading (points)
_CAP_MARGIN_PT    = 4.0    # combined before+after paragraph spacing (points)
# Empirical: Calibri 10pt bold characters per inch (conservative)
_CAP_CHARS_PER_IN = 11.5

# ── Style ─────────────────────────────────────────────────────────────────────
_NAVY   = RGBColor(0x1F, 0x49, 0x7D)
_BLACK  = RGBColor(0x1A, 0x1A, 0x1A)
_GRAY   = RGBColor(0x55, 0x55, 0x55)
_GREEN  = RGBColor(0x1A, 0xB7, 0x38)
_FONT   = "Calibri"

_LOGO_PATH = Path(__file__).parent.parent.parent / "assets" / "sterling_logo.png"
_LOGO_W    = 2.60   # inches  (+30% from original 2.00")


def _init_geometry():
    """
    Compute _MARGIN_T and _BODY_H from the actual logo dimensions so the
    native Word header never overflows into the body regardless of logo size.

    Formula:
      header_distance (0.25") + logo_height + svc_para (~0.19") + rule_para
      (~0.21") + safety_buffer (0.25")
    """
    global _MARGIN_T, _BODY_H
    logo_h = 0.65  # fallback if logo unreadable
    try:
        img = PILImage.open(_LOGO_PATH)
        img = ImageOps.exif_transpose(img)
        w_px, h_px = img.size
        if w_px and h_px:
            logo_h = _LOGO_W * h_px / w_px
    except Exception:
        pass
    # header_distance + logo + service-photos para + rule para (incl. default
    # line height for empty paragraph) + 0.25" safety buffer
    svc_para_h  = (3 + _CAP_FONT_PT + 2) / 72   # before + font + after (pt→in)
    rule_para_h = (12 + 3) / 72                  # default line height + after
    hdr_total   = 0.25 + logo_h + svc_para_h + rule_para_h
    _MARGIN_T   = round(hdr_total + 0.25, 2)
    _BODY_H     = round(_PAGE_H - _MARGIN_T - _MARGIN_B, 3)


_init_geometry()


# ── Layout registry ───────────────────────────────────────────────────────────
# Key = user-facing code, values drive grid math
_LAYOUTS = {
    "3x2":  {"cols": 2, "rows": 3, "per_page": 6},
    "3x3":  {"cols": 3, "rows": 3, "per_page": 9},
    "2x2":  {"cols": 2, "rows": 2, "per_page": 4},
    "full": {"cols": 1, "rows": 1, "per_page": 1},
}

_XML_SPACE = "{http://www.w3.org/XML/1998/namespace}space"


# ═══════════════════════════════════════════════════════════════════════════════
# XML / python-docx helpers
# ═══════════════════════════════════════════════════════════════════════════════

def _ensure_pPr(para):
    """Return the <w:pPr> element for a paragraph, creating it if absent."""
    pPr = para._p.find(qn("w:pPr"))
    if pPr is None:
        pPr = OxmlElement("w:pPr")
        para._p.insert(0, pPr)
    return pPr


def _set_para_spacing(para, before_pt: float = 0, after_pt: float = 0,
                      line_pt: float = None):
    pPr = _ensure_pPr(para)
    for el in pPr.findall(qn("w:spacing")):
        pPr.remove(el)
    sp = OxmlElement("w:spacing")
    sp.set(qn("w:before"), str(int(before_pt * 20)))
    sp.set(qn("w:after"),  str(int(after_pt  * 20)))
    if line_pt is not None:
        sp.set(qn("w:line"),     str(int(line_pt * 20)))
        sp.set(qn("w:lineRule"), "exact")
    pPr.append(sp)


def _add_run(para, text: str, bold=False, size_pt: float = 9,
             color: RGBColor = None, font: str = _FONT):
    run = para.add_run(text)
    run.bold = bold
    run.font.name = font
    run.font.size = Pt(size_pt)
    if color:
        run.font.color.rgb = color
    return run


def _add_field_run(para, instruction: str, size_pt: float = 10):
    """Insert a Word field code (PAGE, NUMPAGES …) into *para*."""
    # begin
    r_begin = para.add_run()
    r_begin.font.size = Pt(size_pt)
    r_begin.font.name = _FONT
    r_begin.font.color.rgb = _GRAY
    fc = OxmlElement("w:fldChar")
    fc.set(qn("w:fldCharType"), "begin")
    r_begin._r.append(fc)

    # instruction
    r_instr = para.add_run()
    r_instr.font.size = Pt(size_pt)
    r_instr.font.name = _FONT
    r_instr.font.color.rgb = _GRAY
    it = OxmlElement("w:instrText")
    it.set(_XML_SPACE, "preserve")
    it.text = f" {instruction} "
    r_instr._r.append(it)

    # separate (provides display text before update)
    r_sep = para.add_run()
    r_sep.font.size = Pt(size_pt)
    r_sep.font.name = _FONT
    r_sep.font.color.rgb = _GRAY
    fc2 = OxmlElement("w:fldChar")
    fc2.set(qn("w:fldCharType"), "separate")
    r_sep._r.append(fc2)

    # placeholder display text
    r_disp = para.add_run("0")
    r_disp.font.size = Pt(size_pt)
    r_disp.font.name = _FONT
    r_disp.font.color.rgb = _GRAY

    # end
    r_end = para.add_run()
    r_end.font.size = Pt(size_pt)
    r_end.font.name = _FONT
    r_end.font.color.rgb = _GRAY
    fc3 = OxmlElement("w:fldChar")
    fc3.set(qn("w:fldCharType"), "end")
    r_end._r.append(fc3)


def _remove_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    for el in tblPr.findall(qn("w:tblBorders")):
        tblPr.remove(el)
    bdr = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{side}")
        e.set(qn("w:val"),   "none")
        e.set(qn("w:sz"),    "0")
        e.set(qn("w:space"), "0")
        e.set(qn("w:color"), "auto")
        bdr.append(e)
    tblPr.append(bdr)


def _set_table_width(table, width_in: float):
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    for el in tblPr.findall(qn("w:tblW")):
        tblPr.remove(el)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"),    str(int(width_in * 1440)))
    tblW.set(qn("w:type"), "dxa")
    tblPr.append(tblW)


def _set_cell_width(cell, width_in: float):
    tc = cell._tc
    tcPr = tc.find(qn("w:tcPr"))
    if tcPr is None:
        tcPr = OxmlElement("w:tcPr")
        tc.insert(0, tcPr)
    for el in tcPr.findall(qn("w:tcW")):
        tcPr.remove(el)
    tcW = OxmlElement("w:tcW")
    tcW.set(qn("w:w"),    str(int(width_in * 1440)))
    tcW.set(qn("w:type"), "dxa")
    tcPr.append(tcW)


def _zero_cell_margins(cell):
    tc = cell._tc
    tcPr = tc.find(qn("w:tcPr"))
    if tcPr is None:
        tcPr = OxmlElement("w:tcPr")
        tc.insert(0, tcPr)
    for el in tcPr.findall(qn("w:tcMar")):
        tcPr.remove(el)
    mar = OxmlElement("w:tcMar")
    for side in ("top", "left", "bottom", "right"):
        e = OxmlElement(f"w:{side}")
        e.set(qn("w:w"),    "36")   # 0.025"
        e.set(qn("w:type"), "dxa")
        mar.append(e)
    tcPr.append(mar)


def _set_row_height_exact(row, height_in: float):
    """Set an exact (non-growable) row height using <w:trHeight w:hRule='exact'>."""
    tr = row._tr
    trPr = tr.find(qn("w:trPr"))
    if trPr is None:
        trPr = OxmlElement("w:trPr")
        tr.insert(0, trPr)
    for el in trPr.findall(qn("w:trHeight")):
        trPr.remove(el)
    trH = OxmlElement("w:trHeight")
    trH.set(qn("w:val"),   str(int(height_in * 1440)))
    trH.set(qn("w:hRule"), "exact")
    trPr.append(trH)
    # Also mark row as non-splittable across pages
    cant = OxmlElement("w:cantSplit")
    cant.set(qn("w:val"), "1")
    trPr.append(cant)


def _add_top_border_para(para_or_container, add_to=None, color="1A1A1A",
                          sz="6", space="4"):
    """
    Add a top border (horizontal rule) to *para_or_container*.
    If add_to is a Document/header/footer, add a new paragraph there and apply.
    Returns the paragraph.
    """
    if add_to is not None:
        p = add_to.add_paragraph()
    else:
        p = para_or_container
    _set_para_spacing(p, before_pt=0, after_pt=0)
    pPr = _ensure_pPr(p)
    pBdr = OxmlElement("w:pBdr")
    bdr = OxmlElement("w:top")
    bdr.set(qn("w:val"),   "single")
    bdr.set(qn("w:sz"),    sz)
    bdr.set(qn("w:space"), space)
    bdr.set(qn("w:color"), color)
    pBdr.append(bdr)
    pPr.append(pBdr)
    return p


def _add_bottom_border_para(container, before_pt=0, after_pt=4,
                             color="1A1A1A", sz="6", space="1"):
    """Add a paragraph with a bottom border (horizontal rule) to *container*."""
    p = container.add_paragraph()
    _set_para_spacing(p, before_pt=before_pt, after_pt=after_pt)
    pPr = _ensure_pPr(p)
    pBdr = OxmlElement("w:pBdr")
    bdr = OxmlElement("w:bottom")
    bdr.set(qn("w:val"),   "single")
    bdr.set(qn("w:sz"),    sz)
    bdr.set(qn("w:space"), space)
    bdr.set(qn("w:color"), color)
    pBdr.append(bdr)
    pPr.append(pBdr)
    return p


# ═══════════════════════════════════════════════════════════════════════════════
# Page setup
# ═══════════════════════════════════════════════════════════════════════════════

def _configure_page(doc):
    section = doc.sections[0]
    section.page_width      = Inches(_PAGE_W)
    section.page_height     = Inches(_PAGE_H)
    section.left_margin     = Inches(_MARGIN_L)
    section.right_margin    = Inches(_MARGIN_R)
    section.top_margin      = Inches(_MARGIN_T)
    section.bottom_margin   = Inches(_MARGIN_B)
    section.header_distance = Inches(0.25)
    section.footer_distance = Inches(0.25)


# ═══════════════════════════════════════════════════════════════════════════════
# Native Word header (logo + service-photos line + rule) — auto-repeats
# ═══════════════════════════════════════════════════════════════════════════════

def _build_header(doc, report_date: str, prepared_by: str):
    """
    Populate the section's native Word header.  Content:
      [Logo 2"  |  date / prepared_by (right-aligned)]
      Service photos are provided below.
      ─────────────────────────────────── (black bottom-border rule)
    """
    section = doc.sections[0]
    header  = section.header
    header.is_linked_to_previous = False

    # Clear any default paragraph Word inserts into the header
    for p in list(header.paragraphs):
        elem = p._p
        if elem.getparent() is not None:
            elem.getparent().remove(elem)

    # ── Logo + info table ─────────────────────────────────────────────────────
    tbl = header.add_table(rows=1, cols=2, width=Inches(_BODY_W))
    _remove_table_borders(tbl)
    _set_table_width(tbl, _BODY_W)

    logo_col_w = _LOGO_W + 0.15
    info_col_w = _BODY_W - logo_col_w

    logo_cell = tbl.rows[0].cells[0]
    info_cell = tbl.rows[0].cells[1]
    _set_cell_width(logo_cell, logo_col_w)
    _set_cell_width(info_cell, info_col_w)
    _zero_cell_margins(logo_cell)
    _zero_cell_margins(info_cell)

    # Logo image (or text fallback)
    logo_p = logo_cell.paragraphs[0]
    logo_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _set_para_spacing(logo_p, before_pt=0, after_pt=0)
    if _LOGO_PATH.exists():
        logo_p.add_run().add_picture(str(_LOGO_PATH), width=Inches(_LOGO_W))
    else:
        _add_run(logo_p, "STERLING STORMWATER", bold=True, size_pt=11, color=_GREEN)

    # Right-side: date / prepared-by
    info_p = info_cell.paragraphs[0]
    info_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _set_para_spacing(info_p, before_pt=4, after_pt=0)
    lines = []
    if report_date:
        lines.append((report_date, False, 8.5))
    if prepared_by:
        lines.append((f"Prepared by: {prepared_by}", False, 8.0))
    for i, (text, bold, size) in enumerate(lines):
        if i > 0:
            info_p.add_run().add_break()
        _add_run(info_p, text, bold=bold, size_pt=size, color=_BLACK)

    # ── "Service photos are provided below." ─────────────────────────────────
    svc_p = header.add_paragraph()
    _set_para_spacing(svc_p, before_pt=2, after_pt=2)
    _add_run(svc_p, "Service photos are provided below.",
             bold=False, size_pt=_CAP_FONT_PT, color=_BLACK)

    # ── Black horizontal rule ─────────────────────────────────────────────────
    # The rule is a bottom-border on an empty paragraph.  Without an explicit
    # exact line height, Word renders it at the default style height (~12pt),
    # which appears as a visible blank line above the rule.  Set line_pt=1 so
    # the paragraph collapses to 1pt — the border still renders, gap disappears.
    rule_p = _add_bottom_border_para(header, before_pt=0, after_pt=0,
                                     color="1A1A1A", sz="6", space="1")
    _set_para_spacing(rule_p, before_pt=0, after_pt=0, line_pt=1.0)


# ═══════════════════════════════════════════════════════════════════════════════
# Native Word footer (separator + contact + page numbers)
# ═══════════════════════════════════════════════════════════════════════════════

def _build_footer(doc):
    """
    Footer layout:
      ─── (top-border on first paragraph = separator rule)
      Please contact … (10pt centered)
      Page N of M     (10pt centered)
    """
    _FPT = 10   # footer font size

    section = doc.sections[0]
    footer  = section.footer
    footer.is_linked_to_previous = False

    for p in list(footer.paragraphs):
        elem = p._p
        if elem.getparent() is not None:
            elem.getparent().remove(elem)

    # ── Contact line — the black rule is its TOP border (no separate paragraph)
    # Putting the border on the text paragraph eliminates the blank-line gap
    # that an empty rule paragraph would create.
    contact_p = footer.add_paragraph()
    contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_para_spacing(contact_p, before_pt=0, after_pt=2)
    pPr = _ensure_pPr(contact_p)
    pBdr = OxmlElement("w:pBdr")
    top_bdr = OxmlElement("w:top")
    top_bdr.set(qn("w:val"),   "single")
    top_bdr.set(qn("w:sz"),    "6")
    top_bdr.set(qn("w:space"), "4")   # 4 pt between rule and text
    top_bdr.set(qn("w:color"), "1A1A1A")
    pBdr.append(top_bdr)
    pPr.append(pBdr)
    _add_run(contact_p,
             "Please contact support@sterlingstormwater.com with questions or "
             "comments regarding the information provided.",
             bold=False, size_pt=_FPT, color=_GRAY)

    # Page N of M
    page_p = footer.add_paragraph()
    page_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_para_spacing(page_p, before_pt=0, after_pt=0)
    _add_run(page_p, "Page ", size_pt=_FPT, color=_GRAY)
    _add_field_run(page_p, "PAGE",     size_pt=_FPT)
    _add_run(page_p, " of ",  size_pt=_FPT, color=_GRAY)
    _add_field_run(page_p, "NUMPAGES", size_pt=_FPT)


# ═══════════════════════════════════════════════════════════════════════════════
# Image helpers
# ═══════════════════════════════════════════════════════════════════════════════

def _get_display_size(filepath: str, max_w: float, max_h: float):
    """Return (w_in, h_in) fitted inside (max_w, max_h), EXIF-corrected."""
    try:
        img = PILImage.open(filepath)
        img = ImageOps.exif_transpose(img)
        w_px, h_px = img.size
        if not w_px or not h_px:
            return max_w, max_h
        scale = min(max_w / w_px, max_h / h_px)
        return w_px * scale, h_px * scale
    except Exception:
        return max_w, max_h


# ═══════════════════════════════════════════════════════════════════════════════
# Dynamic caption height
# ═══════════════════════════════════════════════════════════════════════════════

def _caption_height_for(captions: list, cell_width_in: float) -> float:
    """
    Return the height (inches) needed to render the *longest* caption in
    *captions* at _CAP_FONT_PT inside a cell of *cell_width_in* inches.

    Strategy:
      1. Estimate characters per line from cell width × _CAP_CHARS_PER_IN.
      2. Find worst-case line count across all captions (include order-number
         prefix "(99)  " ≈ 6 chars in the estimate).
      3. Convert line count → inches using _CAP_LINE_H_PT + _CAP_MARGIN_PT.
    Minimum return value: one line + margins (never zero).
    """
    chars_per_line = max(1, int(cell_width_in * _CAP_CHARS_PER_IN))
    max_lines = 1
    for cap in captions:
        text = str(cap).strip()
        # Prefix "(N)  " adds ~6 chars; use a worst-case 3-digit order number
        full_text = f"(999)  {text}" if text else "(999)"
        lines = math.ceil(len(full_text) / chars_per_line)
        max_lines = max(max_lines, lines)
    height_pt = max_lines * _CAP_LINE_H_PT + _CAP_MARGIN_PT
    return height_pt / 72.0


# ═══════════════════════════════════════════════════════════════════════════════
# Single continuous photo table  (no manual page breaks between photo groups)
# ═══════════════════════════════════════════════════════════════════════════════
#
# Why a SINGLE table instead of one-table-per-page with page breaks?
#
# The exact-height rows sum to (body_h - ROW_SAFETY), leaving only ~10pt of
# slack on each page.  A page-break paragraph has a default line height of
# ~12pt.  Because 12pt > 10pt slack, the paragraph overflows to the NEXT page
# alone, creating a blank page before every photo section.
#
# With a single continuous table and cantSplit rows, Word paginates naturally
# at every `rows` row boundary — no paragraph overhead, no blank pages.

def _build_photo_tables(doc, sorted_photos: list, layout_cfg: dict,
                         include_date: bool, global_photo_date: str):
    """
    Build ONE table containing all photos.

    Row heights are set to `exact` = (body_h − safety) / rows_per_page so that
    exactly `rows` rows fill one page.  `cantSplit` keeps each row whole.
    Word paginates the table naturally — no manual page-break paragraphs needed.

    Padding: only the last partial row is padded to `cols` (never a full page),
    so the last page has empty cells but never an entirely blank extra page.
    """
    cols     = layout_cfg["cols"]
    rows     = layout_cfg["rows"]

    row_h   = (_BODY_H - _ROW_SAFETY) / rows
    col_w   = _BODY_W / cols
    photo_w = col_w - 2 * _CELL_PAD

    # Dynamic caption height — pre-scan all captions for worst-case line count
    all_captions = [
        (p.caption or "").strip() or f"Photo {p.order}"
        for p in sorted_photos
    ]
    caption_h = _caption_height_for(all_captions, col_w)
    extra_h   = _DATE_H if include_date else 0.0
    photo_h   = max(row_h - caption_h - extra_h - _CELL_PAD, 0.40)

    # Pad to the next multiple of cols only (not per_page) so no blank rows
    padded = list(sorted_photos)
    remainder = len(padded) % cols
    if remainder:
        padded += [None] * (cols - remainder)

    total_rows = len(padded) // cols

    tbl = doc.add_table(rows=total_rows, cols=cols)
    _remove_table_borders(tbl)
    _set_table_width(tbl, _BODY_W)

    for row_idx in range(total_rows):
        row = tbl.rows[row_idx]
        _set_row_height_exact(row, row_h)

        for col_idx in range(cols):
            slot  = row_idx * cols + col_idx
            photo = padded[slot] if slot < len(padded) else None
            cell  = row.cells[col_idx]
            _set_cell_width(cell, col_w)
            _zero_cell_margins(cell)

            if photo is None:
                _set_para_spacing(cell.paragraphs[0], before_pt=0, after_pt=0)
                continue

            # ── Image ─────────────────────────────────────────────────────────
            img_p = cell.paragraphs[0]
            img_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            _set_para_spacing(img_p, before_pt=2, after_pt=2)
            fp = Path(photo.filepath)
            if fp.exists():
                try:
                    w_in, h_in = _get_display_size(str(fp), photo_w, photo_h)
                    img_p.add_run().add_picture(
                        str(fp), width=Inches(w_in), height=Inches(h_in)
                    )
                except Exception:
                    _add_run(img_p, f"[Image error: {photo.filename}]",
                             size_pt=8, color=RGBColor(0xAA, 0, 0))
            else:
                _add_run(img_p, f"[File not found: {photo.filename}]",
                         size_pt=8, color=RGBColor(0xAA, 0, 0))

            # ── Caption ───────────────────────────────────────────────────────
            cap_text = (photo.caption or "").strip() or f"Photo {photo.order}"
            cap_p = cell.add_paragraph()
            cap_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            _set_para_spacing(cap_p, before_pt=1, after_pt=1)
            _add_run(cap_p, f"({photo.order})  {cap_text}",
                     bold=True, size_pt=_CAP_FONT_PT, color=_BLACK)

            # ── Optional date ─────────────────────────────────────────────────
            if include_date:
                photo_date = (getattr(photo, "photo_date", "") or "").strip()
                date_str   = photo_date or global_photo_date
                if date_str:
                    date_p = cell.add_paragraph()
                    date_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    _set_para_spacing(date_p, before_pt=0, after_pt=1)
                    _add_run(date_p, f"Date of Photo: {date_str}",
                             bold=False, size_pt=8, color=_GRAY)


# ═══════════════════════════════════════════════════════════════════════════════
# Notes page
# ═══════════════════════════════════════════════════════════════════════════════

def _force_page_break(doc):
    """
    Insert a page break that is GUARANTEED to land on the current page
    regardless of how little space remains.

    Uses exact line height of 1pt (20 twips) with a 1pt font so Word never
    pushes this paragraph to the next page — it always fits in any remaining
    space, and its embedded page-break run starts the following content on
    the next page cleanly.
    """
    p = doc.add_paragraph()
    _set_para_spacing(p, before_pt=0, after_pt=0, line_pt=1.0)
    run = p.add_run()
    # 1pt font keeps the run's own line-height contribution at minimum
    rPr = OxmlElement("w:rPr")
    sz  = OxmlElement("w:sz")
    sz.set(qn("w:val"), "2")    # 1pt (sz is in half-points)
    rPr.append(sz)
    run._r.insert(0, rPr)
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run._r.append(br)


def _add_notes_page(doc, photos: list):
    """Append a field-notes page after the photo table."""
    _force_page_break(doc)

    hdr_p = doc.add_paragraph()
    hdr_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _set_para_spacing(hdr_p, before_pt=0, after_pt=3)
    _add_run(hdr_p, "FIELD NOTES", bold=True, size_pt=12, color=_BLACK)

    _add_bottom_border_para(doc, before_pt=0, after_pt=6, color="1A1A1A")

    for photo in photos:
        if not photo.notes.strip():
            continue
        entry_p = doc.add_paragraph()
        _set_para_spacing(entry_p, before_pt=3, after_pt=3)
        _add_run(entry_p, f"Photo {photo.order}", bold=True, size_pt=9.5,
                 color=_NAVY)
        sys_str = (photo.system or "").strip()
        if sys_str and sys_str != "Uncategorized":
            _add_run(entry_p, f"  \u2014  {sys_str}", bold=False, size_pt=9,
                     color=_BLACK)
        _add_run(entry_p, ":  ", size_pt=9)
        _add_run(entry_p, photo.notes.strip(), bold=False, size_pt=9,
                 color=_BLACK)


# ═══════════════════════════════════════════════════════════════════════════════
# Terminal paragraph (prevents Word from appending an extra blank page)
# ═══════════════════════════════════════════════════════════════════════════════

def _fix_terminal_paragraph(doc):
    """
    Word requires a paragraph before <w:sectPr>.  If the document ends with a
    table, Word appends its own empty paragraph which can overflow to a new page.
    We insert a guaranteed 1-twip paragraph to own that slot.
    """
    body    = doc.element.body
    sect_pr = body.find(qn("w:sectPr"))
    if sect_pr is None:
        return

    # Remove empty trailing paragraphs (no text, no breaks, no images)
    preceding = sect_pr.getprevious()
    while preceding is not None and preceding.tag == qn("w:p"):
        has_text  = any(t.text and t.text.strip()
                        for t in preceding.iter(qn("w:t")))
        has_break = preceding.find(f".//{qn('w:br')}") is not None
        has_pic   = preceding.find(f".//{qn('w:drawing')}") is not None
        if has_text or has_break or has_pic:
            break
        parent = preceding.getparent()
        prev   = preceding.getprevious()
        parent.remove(preceding)
        preceding = prev

    # Insert a 1-twip terminal paragraph immediately before sectPr
    term = OxmlElement("w:p")
    pPr  = OxmlElement("w:pPr")
    sp   = OxmlElement("w:spacing")
    sp.set(qn("w:before"),   "0")
    sp.set(qn("w:after"),    "0")
    sp.set(qn("w:line"),     "20")      # 1 twip — virtually invisible
    sp.set(qn("w:lineRule"), "exact")
    pPr.append(sp)
    rPr = OxmlElement("w:rPr")
    sz  = OxmlElement("w:sz")
    sz.set(qn("w:val"), "2")            # 1pt — below any visible threshold
    rPr.append(sz)
    pPr.append(rPr)
    term.insert(0, pPr)
    body.insert(list(body).index(sect_pr), term)


# ═══════════════════════════════════════════════════════════════════════════════
# Public entry point
# ═══════════════════════════════════════════════════════════════════════════════

def build_photosheet(
    photos: list,
    output_path: str,
    layout: str = "3x2",
    site_name: str = "",
    report_date: str = "",
    prepared_by: str = "",
    include_date: bool = False,
    global_photo_date: str = "",
) -> str:
    """
    Build and save a Photosheet DOCX.

    Returns the absolute path of the saved file.
    Raises ValueError for empty photo list or unknown layout.
    """
    if not photos:
        raise ValueError("No photos provided.")
    if layout not in _LAYOUTS:
        raise ValueError(
            f"Unknown layout '{layout}'. Choose from: {list(_LAYOUTS)}"
        )

    layout_cfg = _LAYOUTS[layout]

    # Sort by order field, then renumber 1-based with no gaps
    sorted_photos = sorted(photos, key=lambda p: p.order)
    for i, p in enumerate(sorted_photos):
        p.order = i + 1

    # ── Build document ────────────────────────────────────────────────────────
    doc = Document()
    _configure_page(doc)

    # Strip default empty paragraph that Word inserts on Document()
    for para in list(doc.paragraphs):
        elem = para._p
        if elem.getparent() is not None:
            elem.getparent().remove(elem)

    # Native header (repeats on every page automatically)
    _build_header(doc, report_date, prepared_by)

    # Native footer (repeats on every page automatically)
    _build_footer(doc)

    # One table per page with explicit page breaks — native header auto-repeats
    _build_photo_tables(
        doc=doc,
        sorted_photos=sorted_photos,
        layout_cfg=layout_cfg,
        include_date=include_date,
        global_photo_date=global_photo_date,
    )

    # Notes page (only when at least one photo has notes)
    photos_with_notes = [p for p in sorted_photos if p.notes.strip()]
    if photos_with_notes:
        _add_notes_page(doc, sorted_photos)

    # Prevent trailing blank page
    _fix_terminal_paragraph(doc)

    # ── Save ──────────────────────────────────────────────────────────────────
    # Build into memory first so the DOCX is always complete before touching disk.
    buf = io.BytesIO()
    doc.save(buf)
    raw = buf.getvalue()

    out = Path(output_path)
    out.parent.mkdir(parents=True, exist_ok=True)

    # If the target file is open in Word (Windows file lock → PermissionError),
    # fall back to a timestamped filename so the user still gets their file.
    try:
        out.write_bytes(raw)
    except PermissionError:
        ts  = time.strftime("%H%M%S")
        out = out.with_name(f"{out.stem}_{ts}{out.suffix}")
        out.write_bytes(raw)

    return str(out.resolve())
