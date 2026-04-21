"""
app/services/importer.py
Extract text from a previous report (PDF or DOCX).
V1: Basic extraction with simple section heading detection.
V2: Smart NLP-based section parsing and system-level content attribution.
"""

import io
import re
from typing import BinaryIO


# ── Section heading patterns to look for ─────────────────────────────────────
# These are common headings in stormwater inspection reports.
# Add more patterns to improve section detection for your specific report style.

SECTION_PATTERNS = [
    r"summary\s+and\s+findings",
    r"findings\s+and\s+recommendations",
    r"inspection\s+summary",
    r"maintenance\s+summary",
    r"recommendations",
    r"photo\s+documentation",
    r"introduction",
    r"site\s+description",
    r"background",
    r"scope\s+of\s+(work|services)",
    r"bioretention",
    r"underdrain\s+soil\s+filter",
    r"wet\s+pond",
    r"dry\s+pond",
    r"retention\s+pond",
    r"catch\s+basin",
    r"stormwater\s+wetland",
    r"infiltration",
    r"permeable\s+pavement",
]


def extract_text_from_file(uploaded_file) -> dict:
    """
    Main entry point. Detects file type and routes to correct extractor.
    Returns dict with keys: raw_text, sections, page_count, error (if any).
    """
    filename = uploaded_file.name.lower()

    try:
        if filename.endswith(".pdf"):
            return _extract_from_pdf(uploaded_file)
        elif filename.endswith(".docx"):
            return _extract_from_docx(uploaded_file)
        else:
            return {"error": f"Unsupported file type: {filename}"}
    except Exception as e:
        return {"error": str(e), "raw_text": "", "sections": {}, "page_count": 0}


def _extract_from_pdf(uploaded_file) -> dict:
    """Extract text from PDF using pdfplumber."""
    try:
        import pdfplumber
    except ImportError:
        return {"error": "pdfplumber not installed. Run: pip install pdfplumber"}

    file_bytes = uploaded_file.read()
    raw_text = ""
    page_count = 0

    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        page_count = len(pdf.pages)
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                raw_text += text + "\n\n"

    sections = _parse_sections(raw_text)
    return {
        "raw_text": raw_text,
        "sections": sections,
        "page_count": page_count,
    }


def _extract_from_docx(uploaded_file) -> dict:
    """Extract text from DOCX using python-docx — paragraphs AND table cells.

    Sterling cover pages are built entirely as Word tables, so paragraph-only
    extraction misses every header field (site name, date, inspector, etc.).
    This version iterates both doc.paragraphs and every cell in doc.tables,
    deduplicating so shared text (e.g. a cell that is also a paragraph) is not
    doubled.
    """
    try:
        from docx import Document
    except ImportError:
        return {"error": "python-docx not installed. Run: pip install python-docx"}

    file_bytes = uploaded_file.read()
    doc = Document(io.BytesIO(file_bytes))

    seen: set = set()
    paragraphs: list[str] = []

    def _add(text: str) -> None:
        t = text.strip()
        if t and t not in seen:
            seen.add(t)
            paragraphs.append(t)

    # ── Body paragraphs (non-table content) ──────────────────────────────────
    for para in doc.paragraphs:
        _add(para.text)

    # ── Table cells — critical for Sterling cover page layout ─────────────────
    # Each cell may contain multiple paragraphs; collect them individually so
    # adjacent cells don't run together (preserves field: value separation).
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _add(para.text)

    raw_text = "\n\n".join(paragraphs)
    sections = _parse_sections(raw_text)

    return {
        "raw_text": raw_text,
        "sections": sections,
        "page_count": "N/A (DOCX)",
    }


def _parse_sections(raw_text: str) -> dict:
    """
    Split raw extracted text into named sections based on heading patterns.
    V1: Simple regex-based line scan. Not perfect — designed to get 80% coverage.
    V2: Use spaCy or heading style metadata for more accurate parsing.
    """
    if not raw_text:
        return {}

    lines = raw_text.split("\n")
    sections = {}
    current_section = "Introduction / General"
    buffer = []

    for line in lines:
        stripped = line.strip()
        if not stripped:
            buffer.append("")
            continue

        # Check if this line matches a section heading
        matched_heading = _detect_heading(stripped)
        if matched_heading and len(stripped) < 120:  # headings are short
            # Save previous section
            content = "\n".join(buffer).strip()
            if content:
                sections[current_section] = content
            current_section = matched_heading
            buffer = []
        else:
            buffer.append(line)

    # Save last section
    content = "\n".join(buffer).strip()
    if content:
        sections[current_section] = content

    # Remove empty sections
    sections = {k: v for k, v in sections.items() if v.strip()}

    return sections


def extract_fields(raw_text: str) -> dict:
    """
    Extract structured field values from raw text of a prior Sterling report.
    Used to auto-fill the setup form on import.

    Fix #7: All field patterns now use look-ahead terminators so that adjacent
    PDF table-cell text (which pdfplumber often concatenates on the same line)
    does not bleed into the wrong field.  Fields are left blank when confidence
    is low — no guessing.

    Returns dict with keys:
      site_name, site_address, prepared_by, inspection_date,
      report_type, system_types (list), raw_summary (str)
    All string fields default to "" if not found; system_types defaults to [].
    """
    result = {
        "site_name": "",
        "site_address": "",
        "prepared_by": "",
        "inspection_date": "",
        "report_type": "",
        "system_types": [],
        "raw_summary": "",
    }

    if not raw_text:
        return result

    # ── Shared look-ahead guard ───────────────────────────────────────────────
    # Stops extraction the moment any of these known next-field labels appear,
    # preventing adjacent PDF table cells from contaminating the current field.
    _NEXT_FIELD = (
        r"(?=\s*(?:"
        r"Number\s*(?:&|and)?\s*Type\b"          # "Number & Type of Stormwater…"
        r"|Inspection\s+Performed\s+By\b"
        r"|Number\s+of\s+Pages\b"
        r"|Report\s+Prepared\s+By\b"
        r"|Inspection\s+Company\b"
        r"|Stormwater\s+Compliance"
        r"|\n"                                    # any line break also stops it
        r"|\Z"
        r"))"
    )

    # ── report_type ──────────────────────────────────────────────────────────
    if re.search(r"inspection\s+and\s+maintenance\s+report", raw_text, re.IGNORECASE):
        result["report_type"] = "Inspection and Maintenance"
    elif re.search(r"maintenance\s+report", raw_text, re.IGNORECASE):
        result["report_type"] = "Maintenance"
    elif re.search(r"inspection\s+report", raw_text, re.IGNORECASE):
        result["report_type"] = "Inspection"

    # ── site_name and site_address ───────────────────────────────────────────
    # Fix #7: terminate at _NEXT_FIELD so "Number & Type…" header never bleeds in.
    m = re.search(
        r"Site\s+Name\s*(?:&|and)?\s*Location\s*:\s*(.+?)" + _NEXT_FIELD,
        raw_text,
        re.IGNORECASE,
    )
    if m:
        full_value = m.group(1).strip().rstrip(",")
        # Pattern 1: "Name - Address" split (explicit dash separator)
        dash_split = re.split(r"\s+-\s+", full_value, maxsplit=1)
        if len(dash_split) == 2:
            result["site_name"]    = dash_split[0].strip()
            result["site_address"] = dash_split[1].strip()
        else:
            # Pattern 2: pdfplumber concatenated adjacent table cells:
            # "Site Name [value]  Address: [value]" on the same line.
            addr_m = re.match(
                r"^(.+?)\s+Address\s*:\s*(.+)$", full_value, re.IGNORECASE
            )
            if addr_m:
                result["site_name"]    = addr_m.group(1).strip()
                result["site_address"] = addr_m.group(2).strip()
            else:
                # Pattern 3: "Name, Address" split (only if name part is short)
                comma_split = full_value.split(",", 1)
                if len(comma_split) == 2 and len(comma_split[0]) < 50:
                    result["site_name"]    = comma_split[0].strip()
                    result["site_address"] = comma_split[1].strip()
                else:
                    result["site_name"] = full_value
    else:
        # Plain "Site Name:" label
        m2 = re.search(
            r"Site\s+Name\s*:\s*(.+?)" + _NEXT_FIELD,
            raw_text, re.IGNORECASE,
        )
        if m2:
            result["site_name"] = m2.group(1).strip()

    # ── prepared_by ──────────────────────────────────────────────────────────
    # Fix #7: terminate at _NEXT_FIELD; also strip trailing parenthetical BMP IDs
    # and credential lines that come from the adjacent "Report Prepared By" cell.
    _PREP_GUARD = (
        r"(?=\s*(?:"
        r"Number\s+of\s+Pages\b"
        r"|Number\s*(?:&|and)?\s*Type\b"
        r"|\(\d+\)\s+[A-Z]"                      # "(2) Bioretention …" pattern
        r"|\n"
        r"|\Z"
        r"))"
    )
    for label_pat in [
        r"Inspection\s+Performed\s+By\s*:\s*(.+?)" + _PREP_GUARD,
        r"Performed\s+By\s*:\s*(.+?)"             + _PREP_GUARD,
        r"Report\s+Prepared\s+By\s*:\s*(.+?)"     + _PREP_GUARD,
        r"Prepared\s+By\s*:\s*(.+?)"              + _PREP_GUARD,
    ]:
        m = re.search(label_pat, raw_text, re.IGNORECASE)
        if m:
            value = m.group(1).strip()
            # Strip trailing parenthetical BMP/permit references  e.g. "(LCWMD BMP# 80-01)"
            value = re.sub(
                r"\s*\((?:LCWMD\s+BMP#?\s*\d+[-\w]*|BMP#?\s*\d+[-\w]*)\)\s*$",
                "", value, flags=re.IGNORECASE,
            ).strip()
            # Strip credential lines appended without newline
            value = re.split(r"\s+Stormwater\s+Compliance", value, flags=re.IGNORECASE)[0].strip()
            if value:
                result["prepared_by"] = value
                break

    # ── inspection_date ──────────────────────────────────────────────────────
    date_patterns = [
        r"performed\s+on\s+(\d{1,2}/\d{1,2}/\d{2,4})",
        r"performed\s+on\s+(\w+\s+\d{1,2},?\s+\d{4})",
        r"inspection\s+(?:was\s+)?performed\s+on\s+(\d{1,2}/\d{1,2}/\d{2,4})",
        r"inspection\s+(?:was\s+)?performed\s+on\s+(\w+\s+\d{1,2},?\s+\d{4})",
        r"inspected\s+on\s+(\d{1,2}/\d{1,2}/\d{2,4})",
        r"inspected\s+on\s+(\w+\s+\d{1,2},?\s+\d{4})",
        r"Inspection\s+Date\s*:\s*(\d{1,2}/\d{1,2}/\d{2,4})",
        r"Inspection\s+Date\s*:\s*(\w+\s+\d{1,2},?\s+\d{4})",
    ]
    for pat in date_patterns:
        m = re.search(pat, raw_text, re.IGNORECASE)
        if m:
            result["inspection_date"] = m.group(1).strip()
            break

    # ── system_types ─────────────────────────────────────────────────────────
    # Fix #7: restrict search to the "Number & Type" section block only.
    systems_found = []
    section_match = re.search(
        r"Number\s*(?:&|and)?\s*Type\s+of\s+Stormwater\s+(?:Components?|BMPs?)"
        r"\s+Inspected\s*:?(.*?)"
        r"(?=\n\n|\n(?:[A-Z][A-Z\s]{3,}:)|\Z)",
        raw_text,
        re.IGNORECASE | re.DOTALL,
    )
    search_block = section_match.group(1).strip() if section_match else ""

    if search_block:
        for m in re.finditer(
            r"\(?\d+\)?\s+([A-Z][A-Za-z\s/\-]+?)(?:\s+\(BMP#?[^)]*\))?\s*(?:\n|$)",
            search_block,
        ):
            candidate = m.group(1).strip()
            # Guard: must be 4–60 chars, not start with a digit, not be a known label
            if (4 <= len(candidate) <= 60
                    and not re.match(r"^\d", candidate)
                    and not re.search(r"pages?\s+in\s+report", candidate, re.IGNORECASE)):
                systems_found.append(candidate)

    result["system_types"] = systems_found

    # ── raw_summary ──────────────────────────────────────────────────────────
    summary_match = re.search(
        r"INSPECTION\s+(?:AND\s+MAINTENANCE\s+)?(?:FINDINGS|SUMMARY)\b(.+?)"
        r"(?:CERTIFICATION|INSPECTION\s+AUTHENTICITY)",
        raw_text,
        re.IGNORECASE | re.DOTALL,
    )
    if not summary_match:
        summary_match = re.search(
            r"MAINTENANCE\s+SUMMARY\b(.+?)"
            r"(?:CERTIFICATION|INSPECTION\s+AUTHENTICITY|$)",
            raw_text,
            re.IGNORECASE | re.DOTALL,
        )
    if summary_match:
        result["raw_summary"] = summary_match.group(1).strip()

    return result


def extract_photo_captions(raw_text: str) -> list[dict]:
    """
    Parse photo captions from Sterling-format report text.

    Sterling caption format (from both DOCX and photosheet):
      (N) SystemType – SystemID – ViewDescription
      (N) SystemType - SystemID - ViewDescription   (ASCII dash variant)

    Returns a list of dicts:
      {order: int, system_type: str, system_id: str, view: str}

    Only entries where system_id looks like a real ID (letters + digits + hyphens,
    2–12 chars) are included.  Plain sentences that happen to start with "(N)" are
    filtered out by the system_id validation.
    """
    if not raw_text:
        return []

    _SEP = r"\s*[–\-—]\s*"   # en-dash, hyphen, or em-dash with optional whitespace
    _ID  = r"([A-Z]{1,6}-\d{1,4}[A-Z]?)"   # e.g. CB-1, USF-12, BR-2A

    pattern = re.compile(
        r"^\s*\((\d+)\)\s+"    # (N) order
        r"(.+?)"               # system type (lazy — stops at first separator)
        + _SEP +
        _ID                    # system ID
        + r"(?:" + _SEP + r"(.+?))?$",  # optional view description
        re.MULTILINE,
    )

    results: list[dict] = []
    seen_ids: set = set()

    for m in pattern.finditer(raw_text):
        order       = int(m.group(1))
        system_type = m.group(2).strip()
        system_id   = m.group(3).strip()
        view        = (m.group(4) or "").strip()

        # Deduplicate by (system_type, system_id) — same system appears many times
        key = (system_type.lower(), system_id.upper())
        if key in seen_ids:
            continue
        seen_ids.add(key)

        results.append({
            "order":       order,
            "system_type": system_type,
            "system_id":   system_id.upper(),
            "view":        view,
        })

    return results


def _detect_heading(line: str) -> str | None:
    """
    Return a cleaned section heading name if the line looks like a heading,
    otherwise return None.
    """
    line_lower = line.lower().strip()

    # Strip leading numbers, dots, colons (e.g. "1. Summary and Findings")
    cleaned = re.sub(r"^[\d\.\s]+", "", line).strip()
    cleaned = re.sub(r":$", "", cleaned).strip()

    for pattern in SECTION_PATTERNS:
        if re.search(pattern, line_lower):
            return cleaned.title() if cleaned else line.title()

    # Also catch ALL CAPS short lines as potential headings
    if line.isupper() and 3 < len(line) < 80 and not any(c.isdigit() for c in line[:3]):
        return line.title()

    return None
