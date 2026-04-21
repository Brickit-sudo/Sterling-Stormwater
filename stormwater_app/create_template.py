"""
create_template.py
Run once to generate the starter Word template: templates/report_template.docx

Usage:
    python create_template.py

The template defines named styles that the report builder uses.
Customize the template in Word after generating it — logo, colors,
header/footer — without touching any Python code.

This script is safe to re-run. It will overwrite the existing template.
"""

from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── Brand colors — edit these before generating if you have a brand palette ──
PRIMARY_COLOR   = RGBColor(0x1F, 0x49, 0x7D)   # Navy blue
SECONDARY_COLOR = RGBColor(0x26, 0x6B, 0x9F)   # Mid blue
TEXT_COLOR      = RGBColor(0x26, 0x26, 0x26)   # Near-black
ACCENT_COLOR    = RGBColor(0x70, 0xAD, 0x47)   # Green

FONT_BODY    = "Calibri"
FONT_HEADING = "Calibri"


def create_template(output_path: str = "templates/report_template.docx"):
    Path(output_path).parent.mkdir(parents=True, exist_ok=True)

    doc = Document()

    # ── Page setup ────────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin   = Inches(1.0)
        section.right_margin  = Inches(1.0)

    # ── Base Normal style ─────────────────────────────────────────────────────
    normal = doc.styles["Normal"]
    normal.font.name = FONT_BODY
    normal.font.size = Pt(11)
    normal.font.color.rgb = TEXT_COLOR

    # ── Heading 1 ─────────────────────────────────────────────────────────────
    h1 = doc.styles["Heading 1"]
    h1.font.name = FONT_HEADING
    h1.font.size = Pt(16)
    h1.font.bold = True
    h1.font.color.rgb = PRIMARY_COLOR
    h1.paragraph_format.space_before = Pt(12)
    h1.paragraph_format.space_after  = Pt(6)

    # ── Heading 2 ─────────────────────────────────────────────────────────────
    h2 = doc.styles["Heading 2"]
    h2.font.name = FONT_HEADING
    h2.font.size = Pt(13)
    h2.font.bold = True
    h2.font.color.rgb = SECONDARY_COLOR
    h2.paragraph_format.space_before = Pt(10)
    h2.paragraph_format.space_after  = Pt(4)

    # ── Heading 3 ─────────────────────────────────────────────────────────────
    h3 = doc.styles["Heading 3"]
    h3.font.name = FONT_HEADING
    h3.font.size = Pt(11)
    h3.font.bold = True
    h3.font.color.rgb = TEXT_COLOR
    h3.paragraph_format.space_before = Pt(8)
    h3.paragraph_format.space_after  = Pt(2)

    # ── Body text style ───────────────────────────────────────────────────────
    try:
        body_style = doc.styles.add_style("Report Body", 1)  # 1 = paragraph
    except Exception:
        body_style = doc.styles["Normal"]
    body_style.font.name = FONT_BODY
    body_style.font.size = Pt(11)
    body_style.font.color.rgb = TEXT_COLOR
    body_style.paragraph_format.space_after = Pt(4)

    # ── Caption style ─────────────────────────────────────────────────────────
    try:
        cap_style = doc.styles.add_style("Photo Caption", 1)
    except Exception:
        cap_style = doc.styles["Normal"]
    cap_style.font.name = FONT_BODY
    cap_style.font.size = Pt(9)
    cap_style.font.italic = True
    cap_style.font.color.rgb = TEXT_COLOR
    cap_style.paragraph_format.space_after = Pt(12)

    # ── Template instructions page ────────────────────────────────────────────
    # This page is cleared when the report builder loads the template.
    # Its only purpose is to carry the style definitions forward.

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("STORMWATER REPORT TEMPLATE")
    run.font.name = FONT_HEADING
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = PRIMARY_COLOR

    doc.add_paragraph()

    instructions = doc.add_paragraph()
    instructions_text = (
        "This file is the style template for the Stormwater Report Generator.\n\n"
        "HOW TO CUSTOMIZE:\n"
        "1. Open this file in Microsoft Word.\n"
        "2. Add your company logo to the header (Insert → Header → Edit Header → Insert Picture).\n"
        "3. Update header/footer text with your company name, address, and contact info.\n"
        "4. Modify Heading 1, Heading 2, Heading 3 styles in the Styles pane to match your brand.\n"
        "5. Change font colors via the Styles pane (right-click a style → Modify).\n"
        "6. Save this file. The report generator will use your customized styles automatically.\n\n"
        "DO NOT DELETE THIS FILE — the report generator loads it on export.\n\n"
        "The report content on this page is replaced at export time. "
        "Only the styles and header/footer carry through."
    )
    run2 = instructions.add_run(instructions_text)
    run2.font.name = FONT_BODY
    run2.font.size = Pt(10)
    run2.font.color.rgb = TEXT_COLOR

    doc.save(output_path)
    print(f"Template created: {output_path}")
    print("Open it in Word to add your logo and customize header/footer.")


if __name__ == "__main__":
    create_template()
